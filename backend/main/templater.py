from pathlib import Path
from shutil import make_archive
import os
import re
from slugify import slugify
from docx.shared import Pt
import docx
import pandas as pd

ALLOWED_EXTENSIONS = {'docx', 'xlsx'}
ALLOWED_TEMPLATES = {
    'name': '{{наименование}}',
    'inventory_n': '{{инвентарный номер}}',
    'responsible': '{{ответственный сотрудник}}',
    'otss_category': '{{отсс}}',
    'condition': '{{состояние}}',
    'unit_from': '{{откуда поступила}}',
    'in_operation': '{{используется}}',
    'fault_document_requisites': '{{документы о неисправности}}',
    'date_of_receipt': '{{дата поступления}}',
    'number_of_receipt': '{{номер требования}}',
    'requisites': '{{реквизиты книги учета}}',
    'transfer_date': '{{дата передачи}}',
    'otss_requisites': '{{реквизиты отсс}}',
    'spsi_requisites': '{{реквизиты спси}}',
    'transfer_requisites': '{{реквизиты о передаче}}',
    'last_check': '{{последняя проверка}}',
    'comment': '{{примечания}}',
    'user': '{{кому передали}}',
    'components': '{{компоненты}}'
}


def docx_size(filename):
    """
    Осуществляем проверку, не является ли docx файл пустым
    :return: Количество шаблонов в файле
    """
    size = 0
    try:
        doc = docx.Document(filename)
        size = len(find_docx_templates(doc))
    except AttributeError as error:
        print(error)
        # log.exception('Error with {}'.format(str(error)))

    return size


def find_docx_templates(doc):
    """
    Поиск всех шаблонов в docx документе
    :param doc(Document(file)) - открытый docx-файл
    :return: set(шаблонов)
    """
    templates = []
    try:
        for paragraph in doc.paragraphs:
            for match in re.finditer('\\{\\{.*?\\}\\}', paragraph.text):
                templates.insert(0, match.group(0))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    templates = templates + find_docx_templates(cell)
    except AttributeError as error:
        print(error)
        # log.exception('Error with {}'.format(str(error)))

    sorted(list(templates))
    return templates


def get_docx_templates(filename):
    """
    Метод для получения шаблонов из docx файла
    :return: set(шаблонов)
    """
    docx_templates = []
    try:
        doc = docx.Document(filename)
        docx_templates = find_docx_templates(doc)
    except AttributeError as error:
        print(error)
        # log.exception('Error with {}'.format(str(error)))

    return docx_templates


def check_templates(all_templates):
    global ALLOWED_TEMPLATES
    correct_templates = ALLOWED_TEMPLATES.values()
    incorrect_templates = set()
    for template in all_templates:
        if template not in correct_templates:
            incorrect_templates.add(template)
    return incorrect_templates


def prep_data(payload: list):
    global ALLOWED_TEMPLATES
    keys = ALLOWED_TEMPLATES.keys()
    prep_payload = []
    for item in payload:
        prep_item = {}
        for field in item:
            prep_item[ALLOWED_TEMPLATES[field]] = item[field]
            if prep_item[ALLOWED_TEMPLATES[field]] is None:
                prep_item[ALLOWED_TEMPLATES[field]] = ''
        prep_payload += [prep_item]
    return prep_payload


def docx_write(document, substr, replace):
    """
    Основная функция для замены шаблонов из таблицы
    в docx-документ
    :param document: Document(file) - открытый docx-файл
    :param substr: Первая строка в одном из наборов mini_dict
    :param replace: Вторая строка в одном из наборов mini_dict
    :return:
    """
    #style = document.styles['Normal']
    #font = style.font
    #font.name = 'Times New Roman'
    #font.size = Pt(14)

    for parg in document.paragraphs:
        if substr in parg.text:
            inline = parg.runs
            was_replaced = False
            for i in range(len(inline)):
                if substr in inline[i].text:
                    text = inline[i].text.replace(substr, replace)
                    was_replaced = True
                    inline[i].text = text
            if not was_replaced:
                text = parg.text.replace(substr, replace)
                parg.text = text
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_write(cell, substr, replace)


def final_replacement(filename, payload, merge_doc):
    if not os.path.isdir('media/generated'):
        os.mkdir('media/generated')
    final_data = prep_data(payload)
    replaceable_templates = get_docx_templates(filename)
    for item, i in zip(final_data, range(len(final_data))):
        document = docx.Document(filename)
        for template in replaceable_templates:
            try:
                docx_write(document, template, str(item[template]))
                if os.name == 'nt':
                    document.save('/app/media/generated/Документ-{}.docx'.format(i+1))
                elif os.name == 'posix':
                    document.save('/app/media/generated/Документ-{}.docx'.format(i+1))
            except KeyError as error:
                continue
    if merge_doc:
        documents = os.listdir(os.getcwd() + '/generated')
        combine_word_documents(documents)
        result_path = make_archive('/app/media/Документы',
                                   'zip',
                                   root_dir='/app/media/generated',
                                   base_dir='.')
    else:
        result_path = make_archive('/app/media/Документы',
                                   'zip',
                                   root_dir='/app/media/generated',
                                   base_dir='.')
    return result_path


def combine_word_documents(files):
    """
    Метод для слияния файлов в один
    :param files: Пути к документам в папке
    :return:
    """
    merged_document = docx.Document()
    os.chdir(os.getcwd() + '/generated/')
    for index, file in enumerate(files):
        print('index: {}, file: {}'.format(index, str(file)))
        if index == 0:
            merged_document = docx.Document(file)
            merged_document.add_page_break()
        else:
            sub_doc = docx.Document(file)
            if index < len(files) - 1:
                sub_doc.add_page_break()
            for element in sub_doc.element.body:
                merged_document.element.body.append(element)
    # del_all()
    merged_document.save('Документы.docx')
    os.chdir('..')
    return True

