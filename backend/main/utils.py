import os

from PIL import Image
from pylibdmtx.pylibdmtx import encode
from pathlib import Path
from docx import Document
import xlsxwriter


def prepare_data(item):
    for key_i in item:
        if item[key_i] is not None:
            if key_i != 'components':
                item[key_i] = item[key_i].strip()
            else:
                for component in item['components']:
                    for key_c in component:
                        if component[key_c] is not None:
                            if key_c == 'id':
                                continue
                            if key_c != 'location':
                                component[key_c] = component[key_c].strip()
                            else:
                                for key_l in component['location']:
                                    component['location'][key_l] = component['location'][key_l].strip()
    return item


# def create_data_matrix(payload: list):
#     path_to_doc = os.getcwd() + '/media/Коды.docx'
#     row_num = 0
#     for item in payload:
#         data_to_encode = "Inv:" + item['inventory_n'] + \
#                          " Ser:" + item['serial_n']
#         path_to_image = os.getcwd() + '/media/codes/' + data_to_encode + '.png'
#         path_to_resized_image = os.getcwd() + '/media/codes/' + data_to_encode + '_resize.png'
#         encoded = encode(data_to_encode.encode("utf8"))
#         img = Image.frombytes("RGB", (encoded.width, encoded.height), encoded.pixels)
#         img.save(path_to_image)
#         scale_image(path_to_image, path_to_resized_image, 50, 50)
#         if not os.path.isfile(path_to_doc):
#             document = Document()
#             document.save(path_to_doc)
#         document = Document(path_to_doc)
#         tbl = document.add_table(rows=1, cols=1)
#         row_cells = tbl.add_row().cells
#         paragraph = row_cells[0].paragraphs[0]
#         run = paragraph.add_run()
#         run.add_picture(path_to_resized_image)
#         paragraph = row_cells[0].paragraphs[0]
#         run = paragraph.add_run()
#         run.text = data_to_encode
#         document.save(path_to_doc)
#     return path_to_doc


def scale_image(input_image_path,
                output_image_path,
                width=None,
                height=None
                ):
    original_image = Image.open(input_image_path)
    w, h = original_image.size

    if width and height:
        max_size = (width, height)
    elif width:
        max_size = (width, h)
    elif height:
        max_size = (w, height)
    else:
        # No width or height specified
        raise RuntimeError('Width or height required!')

    original_image.thumbnail(max_size, Image.ANTIALIAS)
    original_image.save(output_image_path)

    scaled_image = Image.open(output_image_path)
    width, height = scaled_image.size


def del_all(dir: str, ext: str):
    """
    Удаляет все лишние файлы после завершения работы
    :return:
    """
    for file in Path(os.getcwd() + dir).glob(ext):
        try:
            file.unlink()
        except OSError as error:
            print(error)


def save_image(data_to_encode):
    path_to_image = os.getcwd() + '/media/codes/' + data_to_encode + '.png'
    path_to_resized_image = os.getcwd() + '/media/codes/' + data_to_encode + '_resize.png'
    encoded = encode(data_to_encode.encode("utf8"))
    img = Image.frombytes("RGB", (encoded.width, encoded.height), encoded.pixels)
    img.save(path_to_image)
    scale_image(path_to_image, path_to_resized_image, 60, 60)
    return path_to_resized_image


def get_data_to_encode(item: dict):
    all = []
    data_to_encode = "Inv:" + item['inventory_n'] + \
                     " Ser:" + item['serial_n']
    all += [data_to_encode]
    for c_serial_n in item['c_serial_n']:
        data_to_encode = "Inv:" + item['inventory_n'] + \
                         " Ser:" + c_serial_n
        all += [data_to_encode]
    return all


def create_data_matrix_xlsx(payload: list):
    path_to_doc = 'media/Коды.xlsx'
    rownum = 2
    workbook = xlsxwriter.Workbook(path_to_doc)
    worksheet = workbook.add_worksheet()
    for item in payload:
        data_to_encode = get_data_to_encode(item)
        for row in data_to_encode:
            img = save_image(row)
            worksheet.set_column('A:A', 25)
            worksheet.write('A' + str(rownum), row)
            worksheet.insert_image('B' + str(rownum), img)
            rownum += 3
    workbook.close()
    return path_to_doc
